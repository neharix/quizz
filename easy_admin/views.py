import datetime
import os
import zipfile
from io import BytesIO

import pandas as pd
import pytz
from django.contrib.auth import logout
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.core.paginator import EmptyPage, PageNotAnInteger, Paginator
from django.http import HttpRequest, HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from challenge.models import *

from .containers import ChallengeResult, FQuestion, UserResult
from .forms import AnswerForm, ChallengeForm, QuestionForm


def get_available_dates(challenge: Challenge):
    sessions = TestSession.objects.filter(challenge=challenge)
    dates = []
    for session in sessions:
        dates.append(session.start)

    dates_str = []
    for date in dates:
        if date.strftime("%d.%m.%Y") not in dates_str:
            dates_str.append(date.strftime("%d.%m.%Y"))

    return dates_str


def logout_view(request: HttpRequest):
    logout(request)
    return redirect("home")


def profile_redirect(request: HttpRequest):
    return redirect("home")


def index(request: HttpRequest):
    if request.user.is_authenticated:
        if request.user.is_superuser or request.user.is_staff:
            return redirect("easy_tools")
        else:
            return redirect("/superuser/login/")
    return redirect("/superuser/login/")


def main(request: HttpRequest):
    if request.user.is_superuser:
        users = User.objects.filter(is_superuser=False, is_staff=False)
        context = {"users": users}
        return render(request, "main.html", context)
    return redirect("home")


def challenges(request: HttpRequest):
    challenges = Challenge.objects.all()
    challenge_results = []
    for index in range(len(challenges)):
        challenge_results.append(ChallengeResult(index + 1, challenges[index]))

    return render(request, "challenges.html", {"challenges": challenge_results})


def challenge_result(
    request: HttpRequest, challenge_pk: int, year=None, month=None, day=None
):
    challenge = Challenge.objects.get(pk=challenge_pk)
    if year and month and day:
        sessions = TestSession.objects.filter(
            challenge=challenge, start__year=year, start__month=month, start__day=day
        )
    else:
        sessions = TestSession.objects.filter(challenge=challenge)
    users = [session.user for session in sessions]
    questions = Question.objects.filter(challenge=challenge)

    user_results = []
    pk = 0
    for user in users:
        session = TestSession.objects.get(challenge=challenge, user=user)
        now = datetime.datetime.now(datetime.timezone.utc)

        timezone = pytz.timezone("Asia/Ashgabat")
        if session.end > now.astimezone(timezone):
            is_finished = False
        else:
            is_finished = True

        pk += 1
        user_answers = []
        for question in questions:
            try:
                user_answers.append(
                    UserAnswer.objects.get(user=user, question=question)
                )
            except:
                pass
        user_results.append(
            UserResult(
                pk,
                challenge.pk,
                user,
                user_answers,
                session,
                is_finished,
                len(questions),
            )
        )
    dates = get_available_dates(challenge)

    paginator = Paginator(user_results, 10)
    page_number = request.GET.get("page", 1)
    try:
        user_results = paginator.page(page_number)
    except PageNotAnInteger:
        user_results = paginator.page(1)
    except EmptyPage:
        user_results = paginator.page(paginator.num_pages)

    paginator_range = [i for i in paginator.page_range]

    return render(
        request,
        "challenge_result.html",
        {
            "paginator_range": paginator_range,
            "users": user_results,
            "challenge": challenge,
            "dates": dates,
            "day": day,
            "month": month,
            "year": year,
            "challenge_pk": challenge_pk,
        },
    )


def export_user_result_short(
    request: HttpRequest,
    challenge_id: int,
    user_id: int,
    year=None,
    month=None,
    day=None,
):

    challenge = Challenge.objects.get(pk=challenge_id)
    user = User.objects.get(pk=user_id)
    if year and month and day:
        session = TestSession.objects.get(
            challenge=challenge,
            user=user,
            start__year=year,
            start__month=month,
            start__day=day,
        )
    else:
        session = TestSession.objects.get(challenge=challenge, user=user)

    questions = Question.objects.filter(challenge=challenge)

    user_answers = []
    for question in questions:
        try:
            user_answers.append(UserAnswer.objects.get(user=user, question=question))
        except:
            pass

    if len(user_answers) != 0:
        true_answer = sum([1 if answer.is_true else 0 for answer in user_answers])
        empty_answer = sum([1 if answer.is_empty else 0 for answer in user_answers])
        false_answer = (
            sum([0 if answer.is_true else 1 for answer in user_answers]) - empty_answer
        )
        percent = round(
            (true_answer / (true_answer + false_answer + empty_answer)) * 100
        )
    else:
        true_answer = 0
        empty_answer = 0
        false_answer = 0
        percent = 0

    start = session.start.astimezone(pytz.timezone("Asia/Ashgabat")).strftime(
        "%d-%m-%Y %H:%M:%S"
    )
    end = session.end.astimezone(pytz.timezone("Asia/Ashgabat")).strftime(
        "%d-%m-%Y %H:%M:%S"
    )

    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    head = document.add_heading(level=0)

    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = head.add_run(f'"{challenge.name}" testiň netijeleri')

    run.font.name = "Times New Roman"
    run.font.bold = True
    document.add_paragraph(f"Ady: {user.first_name}")
    document.add_paragraph(f"Familiýasy: {user.last_name}")
    document.add_paragraph(f"Başlan wagty: {start}")
    document.add_paragraph(f"Gutaran wagty: {end}")
    document.add_paragraph(f"Netije: {percent}%")
    document.add_paragraph(f"Sorag sany: {true_answer + false_answer + empty_answer}")
    document.add_paragraph(f"Dogry jogaplaryň sany: {true_answer}")
    document.add_paragraph(f"Boş jogaplaryň sany: {empty_answer}")
    document.add_paragraph(f"Ýalňyş jogaplaryň sany: {false_answer}")

    document.add_paragraph("")
    document.add_paragraph("")

    sign = document.add_paragraph(
        f"Goly:_________________ {user.last_name.capitalize()} {user.first_name.capitalize()}"
    )

    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response["Content-Disposition"] = (
        f'attachment; filename="{user.last_name.capitalize()} {user.first_name.capitalize()}.docx"'
    )

    return response


def export_user_result(
    request: HttpRequest,
    challenge_id: int,
    user_id: int,
    year=None,
    month=None,
    day=None,
):

    challenge = Challenge.objects.get(pk=challenge_id)
    user = User.objects.get(pk=user_id)
    if year and month and day:
        session = TestSession.objects.get(
            challenge=challenge,
            user=user,
            start__year=year,
            start__month=month,
            start__day=day,
        )
    else:
        session = TestSession.objects.get(challenge=challenge, user=user)
    questions = Question.objects.filter(challenge=challenge)

    user_answers = []
    for question in questions:
        try:
            user_answers.append(UserAnswer.objects.get(user=user, question=question))
        except:
            pass

    if len(user_answers) != 0:
        true_answer = sum([1 if answer.is_true else 0 for answer in user_answers])
        empty_answer = sum([1 if answer.is_empty else 0 for answer in user_answers])
        false_answer = sum([0 if answer.is_true else 1 for answer in user_answers])
        percent = round(
            (true_answer / (true_answer + false_answer + empty_answer)) * 100
        )
    else:
        true_answer = 0
        empty_answer = 0
        false_answer = 0
        percent = 0

    start = session.start.astimezone(pytz.timezone("Asia/Ashgabat")).strftime(
        "%d-%m-%Y %H:%M:%S"
    )
    end = session.end.astimezone(pytz.timezone("Asia/Ashgabat")).strftime(
        "%d-%m-%Y %H:%M:%S"
    )

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)

    head = document.add_heading(level=0)

    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = head.add_run(f'"{challenge.name}" testiň netijeleri')
    run.font.name = "Times New Roman"
    run.font.bold = True

    document.add_paragraph(f"Ady: {user.first_name}")
    document.add_paragraph(f"Familiýasy: {user.last_name}")
    document.add_paragraph(f"Başlan wagty: {start}")
    document.add_paragraph(f"Gutaran wagty: {end}")
    document.add_paragraph(f"Netije: {percent}%")
    document.add_paragraph(f"Sorag sany: {true_answer + false_answer + empty_answer}")
    document.add_paragraph(f"Dogry jogaplaryň sany: {true_answer}")
    document.add_paragraph(f"Boş jogaplaryň sany: {empty_answer}")
    document.add_paragraph(f"Ýalňyş jogaplaryň sany: {false_answer}")
    document.add_paragraph(f"")

    pk = 0

    answer_head = document.add_heading(level=0)
    answer_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    answer_run = answer_head.add_run("Jogaplar")
    answer_run.font.name = "Times New Roman"
    answer_run.font.bold = True

    document.add_paragraph(f"")

    for user_answer in user_answers:
        pk += 1
        if user_answer.question.is_image:
            run = document.add_paragraph().add_run(
                f"{pk}) Sorag ID: {user_answer.question.pk}"
            )
        else:
            run = document.add_paragraph().add_run(
                f"{pk}) {user_answer.question.question}"
            )
        if user_answer.is_true:
            run.font.color.rgb = RGBColor(0, 255, 0)
        else:
            if user_answer.is_empty:
                run.font.color.rgb = RGBColor(97, 97, 97)
            else:
                run.font.color.rgb = RGBColor(255, 0, 0)

        try:
            if user_answer.answer.is_image:
                document.add_paragraph(
                    f"Berlen jogap: Jogap ID: {user_answer.answer.pk}"
                )
            else:
                document.add_paragraph(f"Berlen jogap: {user_answer.answer.answer}")
        except AttributeError:
            document.add_paragraph(f"Berlen jogap: Boş")

        true_ans = Answer.objects.get(question=user_answer.question, is_true=True)

        if true_ans.is_image:
            document.add_paragraph(f"Dogry jogap: Jogap ID:{true_ans.pk}")
        else:
            document.add_paragraph(f"Dogry jogap: {true_ans.answer}")

        answered_at = user_answer.answered_at.astimezone(
            pytz.timezone("Asia/Ashgabat")
        ).strftime("%d-%m-%Y %H:%M:%S")
        document.add_paragraph(f"Jogap berlen wagt: {answered_at}")

        document.add_paragraph("")
        document.add_paragraph("")

    sign = document.add_paragraph(
        f"Goly:_________________ {user.last_name.capitalize()} {user.first_name.capitalize()}"
    )

    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response["Content-Disposition"] = (
        f'attachment; filename="{user.last_name.capitalize()} {user.first_name.capitalize()} umumy.docx"'
    )

    return response


def export_all_results(
    request: HttpRequest, challenge_id: int, year=None, month=None, day=None
):
    challenge = Challenge.objects.get(pk=challenge_id)
    if year and month and day:
        sessions = TestSession.objects.filter(
            challenge=challenge, start__year=year, start__month=month, start__day=day
        )
    else:
        sessions = TestSession.objects.filter(challenge=challenge)
    users = [session.user for session in sessions]
    questions = Question.objects.filter(challenge=challenge)

    user_results = []
    pk = 0
    for user in users:
        session = TestSession.objects.get(challenge=challenge, user=user)
        now = datetime.datetime.now(datetime.timezone.utc)

        timezone = pytz.timezone("Asia/Ashgabat")
        if session.end > now.astimezone(timezone):
            is_finished = False
        else:
            is_finished = True

        pk += 1
        user_answers = []
        for question in questions:
            try:
                user_answers.append(
                    UserAnswer.objects.get(user=user, question=question)
                )
            except:
                pass
        user_results.append(
            UserResult(
                pk,
                challenge.pk,
                user,
                user_answers,
                session,
                is_finished,
                len(questions),
            )
        )

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"

    head = document.add_heading(level=0)

    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = head.add_run(f'"{challenge.name}" testiň netijeleri')

    run.font.name = "Times New Roman"
    run.font.bold = True

    table = document.add_table(rows=len(user_results) + 1, cols=8)
    table.style = "Table Grid"

    for index in range(8):
        cell = table.cell(0, index)
        paragraph = cell.paragraphs[0]
        if index == 3:
            cell.width = Inches(2)
        if index == 0:
            run = paragraph.add_run("№")
        elif index == 1:
            run = paragraph.add_run("Familiýasy we Ady")
        elif index == 2:
            run = paragraph.add_run("Başlan wagty")
        elif index == 3:
            run = paragraph.add_run("Tamamlan wagty")
        elif index == 4:
            run = paragraph.add_run("Dogry jogaplar")
        elif index == 5:
            run = paragraph.add_run("Ýalňyş jogaplar")
        elif index == 6:
            run = paragraph.add_run("Boş jogaplar")
        elif index == 7:
            run = paragraph.add_run("Netije")

        run.font.bold = True

    row_id = 1

    for user_result in user_results:
        start = user_result.start.astimezone(pytz.timezone("Asia/Ashgabat")).strftime(
            "%Y-%m-%d %H:%M:%S"
        )
        end = user_result.end.astimezone(pytz.timezone("Asia/Ashgabat")).strftime(
            "%Y-%m-%d %H:%M:%S"
        )

        table.cell(row_id, 0).text = f"{user_result.id}"

        table.cell(row_id, 1).text = f"{user_result.last_name} {user_result.first_name}"
        start_cell = table.cell(row_id, 2)
        start_cell.text = f"{start}"
        start_cell.width = Inches(2)
        table.cell(row_id, 3).text = f"{end}"
        table.cell(row_id, 4).text = f"{user_result.true_answer}"
        table.cell(row_id, 5).text = f"{user_result.false_answer}"
        table.cell(row_id, 6).text = f"{user_result.empty_answer}"
        table.cell(row_id, 7).text = f"{user_result.percent}%"

        row_id += 1

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response["Content-Disposition"] = f'attachment; filename="{challenge.name}.docx"'

    return response


def editable_challenges(request: HttpRequest):
    challenges = Challenge.objects.all()
    challenge_results = []
    for index in range(len(challenges)):
        challenge_results.append(ChallengeResult(index + 1, challenges[index]))

    return render(
        request, "editable_challenges.html", {"challenges": challenge_results}
    )


def delete_challenge(request: HttpRequest, challenge_id: int):
    challenge = Challenge.objects.get(pk=challenge_id)
    challenge.delete()
    return redirect(request.META["HTTP_REFERER"])


def edit_challenge(request: HttpRequest, challenge_id: int):
    challenge = Challenge.objects.get(pk=challenge_id)
    questions = Question.objects.filter(challenge=challenge).order_by("complexity")
    formatted_questions = []
    pk = 1
    for question in questions:
        formatted_questions.append(FQuestion(pk, question))
        pk += 1
    context = {"challenge": challenge, "questions": formatted_questions}
    return render(request, "edit_challenge.html", context)


def add_challenge(request: HttpRequest):
    if request.method == "POST":
        form = ChallengeForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
        return redirect("editable_challenges")
    context = {"form": ChallengeForm()}
    return render(request, "add_challenge.html", context)


def add_question(request: HttpRequest, challenge_id: int):
    challenge = Challenge.objects.get(pk=challenge_id)

    url_path = request.get_full_path()
    url_list = url_path.split("/")
    url_list.remove("add_question")
    url = "/".join(url_list)
    protocol_meta = request.META.get("SERVER_PROTOCOL").split("/")

    if request.method == "POST":
        form = QuestionForm(request.POST, request.FILES)
        if form.is_valid():
            image = form.cleaned_data.get("image")
            print(image)
            data = form.save()
            question = Question.objects.get(pk=data.pk)
            if image is None:
                question.is_image = False
            else:
                question.is_image = True
            question.save()

        return redirect(f"{protocol_meta[0].lower()}://{request.get_host()}{url}")

    context = {"challenge": challenge, "form": QuestionForm()}
    return render(request, "add_question.html", context)


def edit_question(request: HttpRequest, challenge_id: int, question_id: int):
    question = Question.objects.get(pk=question_id)
    if request.method == "POST":
        if request.POST.get("question") != "":
            question.question = request.POST.get("question")
            question.image = None
            question.is_image = False

        else:
            question.image = request.FILES.get("image")
            question.question = None
            question.is_image = True

        question.point = request.POST.get("point")
        complexity = Complexity.objects.get(pk=request.POST.get("complexity"))
        question.complexity = complexity
        question.save()

        url_path = request.get_full_path()
        url_list = url_path.split("/")
        url_list.remove(f"{question_id}")
        url_list.remove(f"edit_question")
        url = "/".join(url_list)
        protocol_meta = request.META.get("SERVER_PROTOCOL").split("/")

        return redirect(f"{protocol_meta[0].lower()}://{request.get_host()}{url}")

    context = {
        "question": question,
        "answers": Answer.objects.filter(question=question),
        "complexities": Complexity.objects.all(),
        "selected_complexity": question.complexity.pk,
    }
    return render(request, "edit_question.html", context)


def delete_question(request: HttpRequest, challenge_id: int, question_id: int):
    question = Question.objects.get(pk=question_id)
    question.delete()
    return redirect(request.META["HTTP_REFERER"])


def add_answer(request: HttpRequest, challenge_id: int, question_id: int):
    question = Question.objects.get(pk=question_id)

    url_path = request.get_full_path()
    url_list = url_path.split("/")
    url_list.remove("add_answer")
    url = "/".join(url_list)
    protocol_meta = request.META.get("SERVER_PROTOCOL").split("/")

    answers_count = len(Answer.objects.filter(question=question))
    if answers_count >= 4:
        return redirect(f"{protocol_meta[0].lower()}://{request.get_host()}{url}")

    if request.method == "POST":
        print(request.POST)
        form = AnswerForm(request.POST, request.FILES)
        if form.is_valid():
            image = form.cleaned_data.get("image")
            data = form.save()
            answer = Answer.objects.get(pk=data.pk)
            if image is None:
                answer.is_image = False
            else:
                answer.is_image = True
            answer.save()

        return redirect(f"{protocol_meta[0].lower()}://{request.get_host()}{url}")

    context = {"question": question, "form": AnswerForm()}
    return render(request, "add_answer.html", context)


def delete_answer(
    request: HttpRequest, challenge_id: int, question_id: int, answer_id: int
):
    answer = Answer.objects.get(pk=answer_id)
    answer.delete()
    return redirect(request.META["HTTP_REFERER"])


def edit_answer(
    request: HttpRequest, challenge_id: int, question_id: int, answer_id: int
):
    answer = Answer.objects.get(pk=answer_id)
    if request.method == "POST":
        print(request.POST)
        if request.POST.get("answer") != "":
            answer.answer = request.POST.get("answer")
            answer.image = None
            answer.is_image = False
        else:
            answer.image = request.FILES.get("image")
            answer.answer = None
            answer.is_image = True
        answer.is_true = (
            True if request.POST.get("is_true").lower() == "true" else False
        )
        answer.save()

        url_path = request.get_full_path()
        url_list = url_path.split("/")
        url_list.remove(f"{answer_id}")
        url_list.remove(f"edit_answer")
        url = "/".join(url_list)
        protocol_meta = request.META.get("SERVER_PROTOCOL").split("/")

        return redirect(f"{protocol_meta[0].lower()}://{request.get_host()}{url}")

    context = {"answer": answer}
    return render(request, "edit_answer.html", context)


def import_from_xlsx(request: HttpRequest, challenge_id: int):
    if request.method == "POST":
        dataframe = pd.read_excel(request.FILES.get("excel"))

        easy_question_count = 0
        medium_question_count = 0
        hard_question_count = 0

        for index in range(len(dataframe["Sorag"])):
            if dataframe["Derejesi"][index] == "Ýeňil":
                easy_question_count += 1
            elif dataframe["Derejesi"][index] == "Ortaça":
                medium_question_count += 1
            elif dataframe["Derejesi"][index] == "Kyn":
                hard_question_count += 1

        default_question_count = min(
            [
                easy_question_count if easy_question_count != 0 else 1,
                medium_question_count if medium_question_count != 0 else 1,
                hard_question_count if hard_question_count != 0 else 1,
            ]
        )

        question_complexity_counts = {
            "Ýeňil": 0,
            "Ortaça": 0,
            "Kyn": 0,
        }

        zip_file_memory = request.FILES.get("zip", None)
        if zip_file_memory is not None:
            zip_file = BytesIO(zip_file_memory.read())
            with zipfile.ZipFile(zip_file, "r") as file:
                images = file.namelist()
        challenge = Challenge.objects.get(pk=challenge_id)

        for index in range(len(dataframe["Sorag"])):
            is_image = False
            question_text = str(dataframe["Sorag"][index])
            if (
                question_text[0:2] == "{{"
                and question_text[len(question_text) - 2 : len(question_text)] == "}}"
            ):
                filename = question_text.split('"')[1]
                if filename in images:
                    is_image = True
                    with zipfile.ZipFile(zip_file, "r") as file:
                        file.extract(filename, f"temp/{filename}")
                    with open(f"temp/{filename}/{filename}", "rb") as file:
                        image = file.read()
                    os.remove(f"temp/{filename}/{filename}")
                    os.rmdir(f"temp/{filename}")

            complexity = Complexity.objects.get(level=dataframe["Derejesi"][index])
            if question_complexity_counts[complexity.level] < default_question_count:
                question_complexity_counts[complexity.level] += 1
                question = (
                    Question.objects.create(
                        question=question_text,
                        challenge=challenge,
                        point=1,
                        complexity=complexity,
                    )
                    if is_image == False
                    else Question.objects.create(
                        challenge=challenge,
                        point=1,
                        complexity=complexity,
                        is_image=is_image,
                        image=ContentFile(image, filename),
                    )
                )
                true_answer = (
                    dataframe["Dogry jogap"][index]
                    if type(dataframe["Dogry jogap"][index]) == int
                    else int(dataframe["Dogry jogap"][index])
                )
                for i in range(1, 5):
                    if type(dataframe[f"{i}-nji jogap"][index]) != float:
                        is_image = False
                        answer_text = str(dataframe[f"{i}-nji jogap"][index])
                        if (
                            answer_text[0:2] == "{{"
                            and answer_text[len(answer_text) - 2 : len(answer_text)]
                            == "}}"
                        ):
                            filename = answer_text.split('"')[1]
                            if filename in images:
                                is_image = True
                                with zipfile.ZipFile(zip_file, "r") as file:
                                    file.extract(filename, f"temp/{filename}")
                                with open(f"temp/{filename}/{filename}", "rb") as file:
                                    image = file.read()
                                os.remove(f"temp/{filename}/{filename}")
                                os.rmdir(f"temp/{filename}")
                        if is_image:
                            answer = (
                                Answer.objects.create(
                                    image=ContentFile(image, filename),
                                    question=question,
                                    is_true=True,
                                    is_image=is_image,
                                )
                                if true_answer == i
                                else Answer.objects.create(
                                    image=ContentFile(image, filename),
                                    question=question,
                                    is_image=is_image,
                                )
                            )
                        else:
                            answer = (
                                Answer.objects.create(
                                    answer=answer_text, question=question, is_true=True
                                )
                                if true_answer == i
                                else Answer.objects.create(
                                    answer=answer_text, question=question
                                )
                            )
            else:
                continue

        return render(
            request,
            "import_from_xlsx.html",
            {"type": "success", "message": "Maglumat gorunda üstünlikli girizildi!"},
        )
    return render(request, "import_from_xlsx.html")


def check_get(request: HttpRequest):
    if request.method == "GET":
        print(request.GET)

    testsession = TestSession.objects.get(pk=54)
    print(testsession.__dict__)
    return HttpResponse(status=200)


def real_time_update(
    request: HttpRequest, challenge_pk: int, year=None, month=None, day=None
):
    challenge = Challenge.objects.get(pk=challenge_pk)

    return render(
        request,
        "online_change_board.html",
        {
            "challenge": challenge,
            "challenge_pk": challenge_pk,
        },
    )


def real_time_chart_update(
    request: HttpRequest, challenge_pk: int, year=None, month=None, day=None
):
    challenge = Challenge.objects.get(pk=challenge_pk)

    return render(
        request,
        "online_chart.html",
        {
            "challenge": challenge,
            "challenge_pk": challenge_pk,
        },
    )


def real_time_challenges(request: HttpRequest):
    challenges = Challenge.objects.all()
    challenge_results = []
    for index in range(len(challenges)):
        challenge_results.append(ChallengeResult(index + 1, challenges[index]))

    return render(
        request, "real_time_challenges.html", {"challenges": challenge_results}
    )
